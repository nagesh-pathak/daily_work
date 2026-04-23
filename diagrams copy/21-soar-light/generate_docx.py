#!/usr/bin/env python3
"""
Generate a professional requirements DOCX from the SOAR-Light README content.
Output: SOAR-Light_C2C_Scheduled_Source_Requirements.docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os
import cairosvg
import tempfile

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

BRAND_BLUE   = RGBColor(0x1B, 0x3A, 0x5C)   # dark navy
ACCENT_BLUE  = RGBColor(0x2E, 0x75, 0xB6)   # softer blue for headings
LIGHT_GRAY   = RGBColor(0xF2, 0xF2, 0xF2)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
BLACK        = RGBColor(0x00, 0x00, 0x00)
CODE_BG      = "E8EDF2"
TABLE_HEADER = "1B3A5C"
TABLE_ALT    = "EAF0F7"


def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, top="single", bottom="single", left="single", right="single",
                     color="BFBFBF", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="{top}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="{bottom}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="{left}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="{right}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcBorders_existing = tcPr.find(qn('w:tcBorders'))
    if tcBorders_existing is not None:
        tcPr.remove(tcBorders_existing)
    tcPr.append(tcBorders)


def styled_table(doc, headers, rows, col_widths=None):
    """Create a professional table with header styling and alternating row colors."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, TABLE_HEADER)
        set_cell_borders(cell, color="1B3A5C")

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.color.rgb = BLACK
            if r_idx % 2 == 1:
                set_cell_shading(cell, TABLE_ALT)
            set_cell_borders(cell, color="D0D0D0")

    # Column widths
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Inches(width)

    doc.add_paragraph("")  # spacing
    return table


def add_code_block(doc, code_text, language=""):
    """Add a formatted code block with background shading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(code_text.rstrip())
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    # shading on the paragraph
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{CODE_BG}"/>')
    pPr.append(shading)


def add_heading_with_color(doc, text, level, color=ACCENT_BLUE):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
    return h


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="BFBFBF"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


# Directory where this script and the SVGs live
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
# Temp directory for PNG conversions
PNG_DIR = os.path.join(SCRIPT_DIR, "_png_cache")
os.makedirs(PNG_DIR, exist_ok=True)


def svg_to_png(svg_filename, dpi=150):
    """Convert an SVG file to a high-res PNG. Returns the PNG path."""
    svg_path = os.path.join(SCRIPT_DIR, svg_filename)
    png_path = os.path.join(PNG_DIR, svg_filename.replace(".svg", ".png"))
    # Scale = dpi / 96 (SVG default is 96 DPI)
    scale = dpi / 96.0
    cairosvg.svg2png(url=svg_path, write_to=png_path, scale=scale)
    return png_path


def add_diagram(doc, svg_filename, caption, width_inches=6.5):
    """Convert SVG → PNG and insert it centered with a caption."""
    png_path = svg_to_png(svg_filename, dpi=200)
    # Add the image centered
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(png_path, width=Inches(width_inches))
    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    cap.paragraph_format.space_after = Pt(12)


# ---------------------------------------------------------------------------
# Document Construction
# ---------------------------------------------------------------------------

doc = Document()

# -- Page setup --
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)

# -- Default font --
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(10)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# Heading styles
for lvl in range(1, 4):
    hs = doc.styles[f"Heading {lvl}"]
    hs.font.name = "Calibri"
    hs.font.color.rgb = ACCENT_BLUE if lvl <= 2 else BRAND_BLUE

# ===================================================================
# COVER PAGE
# ===================================================================
for _ in range(6):
    doc.add_paragraph("")

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("Requirements Document")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = BRAND_BLUE
run.font.name = "Calibri"

doc.add_paragraph("")

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle_p.add_run("Scheduled Source Support for\nSOAR-Light Cloud-to-Cloud (C2C) Flows")
run.font.size = Pt(16)
run.font.color.rgb = ACCENT_BLUE
run.font.name = "Calibri"

doc.add_paragraph("")
add_horizontal_rule(doc)
doc.add_paragraph("")

# Meta info on cover
meta_items = [
    ("Feature", "Enable SOURCE_SCHEDULE for C2C SOAR-Light flows"),
    ("Components", "cdc.flow.api.service, soar-light"),
    ("Branches", "cdc.flow.api.service: c2c-schedule-support-v296\nsoar-light: soar-schedule-support-043"),
    ("Status", "Implementation complete — both components modified and build-verified"),
    ("Date", "April 2026"),
    ("Author", "Nagesh Pathak"),
    ("Classification", "Internal / Engineering"),
]

meta_table = doc.add_table(rows=len(meta_items), cols=2)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (label, value) in enumerate(meta_items):
    cell_l = meta_table.rows[i].cells[0]
    cell_l.text = ""
    cell_l.width = Inches(1.8)
    p = cell_l.paragraphs[0]
    r = p.add_run(label)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = BRAND_BLUE
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    cell_r = meta_table.rows[i].cells[1]
    cell_r.text = ""
    cell_r.width = Inches(4.5)
    p = cell_r.paragraphs[0]
    r = p.add_run(value)
    r.font.size = Pt(10)
    r.font.color.rgb = BLACK

    set_cell_borders(cell_l, top="none", bottom="none", left="none", right="none")
    set_cell_borders(cell_r, top="none", bottom="none", left="none", right="none")

doc.add_page_break()

# ===================================================================
# TABLE OF CONTENTS (manual)
# ===================================================================
add_heading_with_color(doc, "Table of Contents", level=1, color=BRAND_BLUE)
toc_items = [
    "1.  Problem Statement",
    "2.  Current Behavior",
    "3.  Proposed Behavior",
    "4.  Scope",
    "5.  Technical Analysis",
    "6.  Implementation Details",
    "7.  Kubernetes Resource Comparison",
    "8.  End-to-End Flow",
    "9.  Impact Summary",
    "10. Test Validation",
    "11. Resolved Issues",
    "12. Architecture Diagrams",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(1)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = ACCENT_BLUE

doc.add_page_break()

# ===================================================================
# 1. PROBLEM STATEMENT
# ===================================================================
add_heading_with_color(doc, "1. Problem Statement", level=1)

doc.add_paragraph(
    "SOAR-Light supports two script execution modes:"
)
styled_table(doc,
    ["Mode", "Interface", "Trigger", "Source Type"],
    [
        ["Event-driven", "IEventHandler", "Fired per Kafka event", "SOURCE_BLOXONE"],
        ["Scheduled",    "IScheduler",    "Fired on cron schedule", "SOURCE_SCHEDULE"],
    ],
    col_widths=[1.4, 1.4, 1.8, 1.8]
)

doc.add_paragraph(
    'Scheduled source (SOURCE_SCHEDULE) is currently blocked for C2C flows by an explicit '
    'validation check in cdc.flow.api.service. It is only available for on-prem deployments.'
)
doc.add_paragraph(
    'The block exists because the C2C deployment model creates two pods per flow — a grpc-in '
    'pod (reads Kafka events) and a soar-light pod (runs Python). For scheduled flows, the '
    'grpc-in pod has no function since there are no events to consume.'
)

p = doc.add_paragraph()
run = p.add_run(
    "Requirement: Allow SOURCE_SCHEDULE in C2C by creating only the soar-light pod "
    "(skipping grpc-in), eliminating wasted resources."
)
run.bold = True

# Diagram: Event-driven vs Scheduled script types
add_diagram(doc, "event-vs-scheduled.svg",
           "Figure 1 — Event Handler (IEventHandler) vs Scheduled Script (IScheduler)", width_inches=6.3)

add_horizontal_rule(doc)

# ===================================================================
# 2. CURRENT BEHAVIOR
# ===================================================================
add_heading_with_color(doc, "2. Current Behavior", level=1)

add_heading_with_color(doc, "2.1 C2C Flow Resource Creation", level=2)
doc.add_paragraph("For every C2C flow, cdc.flow.api.service currently creates:")
styled_table(doc,
    ["Resource", "Name Pattern", "Purpose"],
    [
        ["ConfigMap",    "cdc-grpc-in-cm-{accountId}-{flowId}",         "grpc-in pod configuration"],
        ["StatefulSet",  "cdc-grpc-in-service-{accountId}-{flowId}",    "grpc-in pod — reads from Kafka"],
        ["ConfigMap",    "cdc-soar-light-cm-{accountId}-{flowId}",      "soar-light pod configuration"],
        ["StatefulSet",  "cdc-soar-light-service-{accountId}-{flowId}", "soar-light pod — runs Python script"],
    ],
    col_widths=[1.2, 3.0, 2.2]
)
p = doc.add_paragraph()
run = p.add_run("Total: 2 pods, 2 ConfigMaps, 2 StatefulSets per flow.")
run.bold = True

add_heading_with_color(doc, "2.2 Validation Block", level=2)
doc.add_paragraph("cdc.flow.api.service/pkg/svc/flow_handler_service.go — validateC2CService():")
add_code_block(doc, '''switch src.Type {
case db.SourceNIOSType, db.SourceIngressType, db.SourceScheduleType:
    return fmt.Errorf("source type %s is not supported for Data Connector in Infoblox Cloud",
        db.SourceTypeToName[src.Type])
}''', "go")
doc.add_paragraph("SOURCE_SCHEDULE is rejected alongside SOURCE_NIOS and SOURCE_INGRESS.")

add_heading_with_color(doc, "2.3 On-Prem vs C2C Comparison", level=2)
styled_table(doc,
    ["Aspect", "On-Prem", "C2C"],
    [
        ["Deployment",        "Single soar-light binary, all flows",  "Per-flow K8s pods"],
        ["Config delivery",   "JSON file on disk",                     "FLOW_JSON env var in ConfigMap"],
        ["Scheduled support", "Yes — internal cron in the binary",     "Blocked at API validation"],
        ["grpc-in pod",       "None",                                  "Required for event-driven flows"],
    ],
    col_widths=[1.5, 2.5, 2.5]
)

add_horizontal_rule(doc)

# ===================================================================
# 3. PROPOSED BEHAVIOR
# ===================================================================
add_heading_with_color(doc, "3. Proposed Behavior", level=1)

add_heading_with_color(doc, "3.1 Scheduled C2C Flow Resources", level=2)
styled_table(doc,
    ["Resource", "Name Pattern", "Created?"],
    [
        ["ConfigMap",    "cdc-grpc-in-cm-{accountId}-{flowId}",         "No"],
        ["StatefulSet",  "cdc-grpc-in-service-{accountId}-{flowId}",    "No"],
        ["ConfigMap",    "cdc-soar-light-cm-{accountId}-{flowId}",      "Yes"],
        ["StatefulSet",  "cdc-soar-light-service-{accountId}-{flowId}", "Yes"],
    ],
    col_widths=[1.2, 3.0, 1.0]
)
p = doc.add_paragraph()
run = p.add_run("Total: 1 pod, 1 ConfigMap, 1 StatefulSet per scheduled flow.")
run.bold = True

add_heading_with_color(doc, "3.2 Scheduled Flow Detection Criteria", level=2)
doc.add_paragraph("A flow is classified as scheduled when:")
add_code_block(doc, 'len(SourceDataTypes) == 0  AND  ScriptSchedule != ""')

add_heading_with_color(doc, "3.3 Start the Cron Scheduler in Cloud Mode", level=2)
doc.add_paragraph(
    "The soar-light binary already contains the cron registration logic in "
    "pkg/dapr/subscriber.go → runDataFlow():"
)
add_code_block(doc, '''if len(flow.DataTypes) == 0 && flow.Schedule != "" {
    cronId, err := s.o.Manager.ProcessScheduledScript(flow)
    s.o.Config.AddCronID(flow.Id, cronId)
    return &flow, nil
}''', "go")
doc.add_paragraph(
    "However, the robfig/cron scheduler was never started in cloud mode (cmd/cloud/main.go). "
    "The on-prem entry point (cmd/onprem/main.go) calls cron.Start(), but the cloud entry "
    "point did not. Without cron.Start(), the cron library registers jobs but its internal "
    "ticker goroutine never fires — so IScheduler.schedule() is never called."
)
p = doc.add_paragraph()
run = p.add_run("Fix: ")
run.bold = True
run = p.add_run("Add cron.Start() + defer cron.Stop() in cmd/cloud/main.go before the ProcessEvent goroutine.")

# Diagram: Full comparison — current vs proposed
doc.add_paragraph("")
add_diagram(doc, "c2c-current-vs-scheduled-comparison.svg",
           "Figure 2 — Current Event-Driven C2C Flow vs Proposed Scheduled C2C Flow (side-by-side)", width_inches=6.5)

add_horizontal_rule(doc)

# ===================================================================
# 4. SCOPE
# ===================================================================
add_heading_with_color(doc, "4. Scope", level=1)

add_heading_with_color(doc, "4.1 In Scope", level=2)
styled_table(doc,
    ["Item", "Repository"],
    [
        ["Remove SOURCE_SCHEDULE from C2C blocked list",       "cdc.flow.api.service"],
        ["Return nil InSvc for scheduled flows",               "cdc.flow.api.service"],
        ["Allow nil InSvc in NewCDCFlowHandler()",             "cdc.flow.api.service"],
        ["Add nil checks in CreateFlow, UpdateFlow, DeleteFlow", "cdc.flow.api.service"],
        ["Fix sync reconciliation loop for nil InSvc",         "cdc.flow.api.service"],
        ["Fix ConfigMap count expectation in sync",            "cdc.flow.api.service"],
        ["Start cron scheduler in cloud mode",                 "soar-light"],
    ],
    col_widths=[4.0, 2.4]
)

add_heading_with_color(doc, "4.2 Out of Scope", level=2)
styled_table(doc,
    ["Item", "Reason"],
    [
        ["Python driver changes",   "Already supports --script-schedule mode"],
        ["CDC API changes",         "No API contract changes needed"],
        ["On-prem flow behavior",   "Unaffected"],
    ],
    col_widths=[2.5, 3.9]
)

add_horizontal_rule(doc)

# ===================================================================
# 5. TECHNICAL ANALYSIS
# ===================================================================
add_heading_with_color(doc, "5. Technical Analysis", level=1)

add_heading_with_color(doc, "5.1 Key Code Path: soar-light Scheduled Execution", level=2)
add_code_block(doc, '''cmd/cloud/main.go
  → viper.GetString("flow.json")           // reads FLOW_JSON from ConfigMap env
  → svc.GetCron().Start()                   // CHANGE 7: start cron scheduler
  → daprSub.ProcessEvent(flowJSON)
    → processEvent()
      → json.Unmarshal → FlowEvent
      → validateConfig(fc)                  // reject on-prem flows
      → validateDestinations(fc.Destinations)
      → runDataFlow(fc)
        → getParsedAndEncodedFlowConfig()
        → if DataTypes==0 && Schedule!="" → ProcessScheduledScript(flow)
          → creates Python venv
          → starts cron job
          → Python calls IScheduler.schedule() on each cron tick''')

add_heading_with_color(doc, "5.2 Key Code Path: cdc.flow.api.service Flow Handler", level=2)
add_code_block(doc, '''FlowEvent received via Dapr PubSub
  → validateC2CService()                   // CHANGE 1: allow SOURCE_SCHEDULE
  → getSrcSvcConf(flowEvent)               // CHANGE 2: return nil for scheduled
  → NewCDCFlowHandler(flowEvent)           // CHANGE 3: accept nil InSvc
    → CDCFlowHandler{InSvc: nil, OutSvc: soarLightSvc}
  → fh.CreateFlow()                        // CHANGE 4: skip nil InSvc operations
    → OutSvc.CreateOrUpdateConfigMap()
    → OutSvc.CreateStatefulSet()''')

add_heading_with_color(doc, "5.3 Sync Reconciliation Loop", level=2)
doc.add_paragraph("syncCDCFlows() runs periodically to reconcile in-cluster state:")
add_code_block(doc, '''syncCDCFlows()
  → for each active flow handler:
    → syncConfigMaps()       // CHANGE 5: handle nil InSvc
    → syncStatefulSets()     // CHANGE 5: handle nil InSvc
    → check ConfigMap count  // CHANGE 6: expect 1 CM for scheduled''')

add_horizontal_rule(doc)

# ===================================================================
# 6. IMPLEMENTATION DETAILS
# ===================================================================
add_heading_with_color(doc, "6. Implementation Details", level=1)

# Change 1
add_heading_with_color(doc, "Change 1: Remove SOURCE_SCHEDULE from blocked list", level=2)
p = doc.add_paragraph()
run = p.add_run("File: ")
run.bold = True
p.add_run("cdc.flow.api.service/pkg/svc/flow_handler_service.go")
add_code_block(doc, ''' switch src.Type {
-case db.SourceNIOSType, db.SourceIngressType, db.SourceScheduleType:
+case db.SourceNIOSType, db.SourceIngressType:
     return fmt.Errorf("source type %s is not supported ...",
         db.SourceTypeToName[src.Type])
 }''')

# Change 2
add_heading_with_color(doc, "Change 2: Return nil InSvc for scheduled sources", level=2)
p = doc.add_paragraph()
run = p.add_run("File: ")
run.bold = True
p.add_run("cdc.flow.api.service/pkg/svc/utils.go — getSrcSvcConf()")
add_code_block(doc, ''' func getSrcSvcConf(config *dapr.FlowEvent) cdcSvc {
+    if len(config.Flow.SourceDataTypes) == 0 && config.Flow.ScriptSchedule != "" {
+        return nil
+    }
     return NewGrpcInSvcCli(config)
 }''')

# Change 3
add_heading_with_color(doc, "Change 3: Allow nil InSvc in NewCDCFlowHandler()", level=2)
p = doc.add_paragraph()
run = p.add_run("File: ")
run.bold = True
p.add_run("cdc.flow.api.service/pkg/svc/flow_handler.go")
add_code_block(doc, ''' inSvc := getSrcSvcConf(flowEvent)
-if inSvc == nil {
+if inSvc == nil && !isScheduledFlow(flowEvent) {
     return nil, fmt.Errorf("invalid source for cloud flow, flow: %v", flowEvent.Flow.Id)
 }''')
doc.add_paragraph("New helper function:")
add_code_block(doc, '''func isScheduledFlow(fe *dapr.FlowEvent) bool {
    return len(fe.Flow.SourceDataTypes) == 0 && fe.Flow.ScriptSchedule != ""
}''')

# Change 4
add_heading_with_color(doc, "Change 4: Nil-safe InSvc in CreateFlow(), UpdateFlow(), DeleteFlow()", level=2)
p = doc.add_paragraph()
run = p.add_run("File: ")
run.bold = True
p.add_run("cdc.flow.api.service/pkg/svc/flow_handler.go")
doc.add_paragraph("All three methods wrap InSvc calls with nil checks:")
add_code_block(doc, '''if fh.InSvc != nil {
    if err := fh.InSvc.CreateOrUpdateConfigMap(); err != nil { return err }
    if err := fh.InSvc.CreateStatefulSet(); err != nil { return err }
}
// OutSvc always executes
if err := fh.OutSvc.CreateOrUpdateConfigMap(); err != nil { return err }
if err := fh.OutSvc.CreateStatefulSet(); err != nil { return err }''')
doc.add_paragraph("Same pattern applied to UpdateFlow() and DeleteFlow().")

# Change 5
add_heading_with_color(doc, "Change 5: Fix sync functions for nil InSvc", level=2)
p = doc.add_paragraph()
run = p.add_run("File: ")
run.bold = True
p.add_run("cdc.flow.api.service/pkg/svc/utils.go")

p = doc.add_paragraph()
run = p.add_run("syncConfigMaps(): ")
run.bold = True
p.add_run("Declare var inSvcCM *corev1.ConfigMap, populate only if flh.InSvc != nil. "
          "Add nil check when matching ConfigMap names.")

p = doc.add_paragraph()
run = p.add_run("syncStatefulSets(): ")
run.bold = True
p.add_run("Remove flh.InSvc == nil from the error condition. Conditionally prepare "
          "and sync InSvc StatefulSet only when non-nil.")

# Change 6
add_heading_with_color(doc, "Change 6: Fix ConfigMap count expectation in syncCDCFlows()", level=2)
p = doc.add_paragraph()
run = p.add_run("File: ")
run.bold = True
p.add_run("cdc.flow.api.service/pkg/svc/utils.go")
add_code_block(doc, '''-if len(cms) != 2 {
-    logrus.Warnf("Unexpected number of ConfigMaps (%d) ... expected 2, ...")
+expectedCmCount := 2
+if flh.InSvc == nil {
+    expectedCmCount = 1
+}
+if len(cms) != expectedCmCount {
+    logrus.Warnf("Unexpected number of ConfigMaps (%d) ... expected %d, ...",
+        len(cms), flowId, expectedCmCount, names)
 }''')

# Change 7
add_heading_with_color(doc, "Change 7: Start cron scheduler in cloud mode", level=2)
p = doc.add_paragraph()
run = p.add_run("File: ")
run.bold = True
p.add_run("soar-light/cmd/cloud/main.go")
doc.add_paragraph(
    "The on-prem entry point (cmd/onprem/main.go, ~line 71) calls svc.GetCron().Start(), "
    "but the cloud entry point was missing this call entirely. Without it, robfig/cron "
    "registers jobs but the internal ticker goroutine never starts."
)
add_code_block(doc, '''  flowJSON := viper.GetString("flow.json")
  logger.Debugf("flowJSON config: %v", flowJSON)

+ // Start the cron scheduler so scheduled flows can fire.
+ cronScheduler := svc.GetCron()
+ cronScheduler.Start()
+ defer cronScheduler.Stop()
+
  go func() {
      if err = daprSub.ProcessEvent(flowJSON); err != nil {
          doneC <- err
      }
  }()''')

add_horizontal_rule(doc)

# ===================================================================
# 7. KUBERNETES RESOURCE COMPARISON
# ===================================================================
add_heading_with_color(doc, "7. Kubernetes Resource Comparison", level=1)

add_heading_with_color(doc, "Event-Driven Flow (unchanged)", level=2)
add_code_block(doc, '''$ kubectl get all -n cdc-data-flow | grep {flowId}
pod/cdc-grpc-in-service-{acctId}-{flowId}-0          1/1   Running
pod/cdc-soar-light-service-{acctId}-{flowId}-0       1/1   Running
configmap/cdc-grpc-in-cm-{acctId}-{flowId}
configmap/cdc-soar-light-cm-{acctId}-{flowId}
statefulset.apps/cdc-grpc-in-service-{acctId}-{flowId}
statefulset.apps/cdc-soar-light-service-{acctId}-{flowId}''')

add_heading_with_color(doc, "Scheduled Flow (after changes)", level=2)
add_code_block(doc, '''$ kubectl get all -n cdc-data-flow | grep {flowId}
pod/cdc-soar-light-service-{acctId}-{flowId}-0       1/1   Running
configmap/cdc-soar-light-cm-{acctId}-{flowId}
statefulset.apps/cdc-soar-light-service-{acctId}-{flowId}''')

add_heading_with_color(doc, "Per-Flow Resource Savings", level=2)
styled_table(doc,
    ["Resource", "Event Flow", "Scheduled Flow", "Saved"],
    [
        ["Pods",              "2",      "1",      "1"],
        ["StatefulSets",      "2",      "1",      "1"],
        ["ConfigMaps",        "2",      "1",      "1"],
        ["Memory",            "~200Mi", "~100Mi", "~100Mi"],
        ["CPU",               "~60m",   "~30m",   "~30m"],
        ["Kafka connections", "1",      "0",      "1"],
    ],
    col_widths=[1.8, 1.5, 1.5, 1.5]
)

# Diagram: Scheduled C2C proposed architecture
add_diagram(doc, "scheduled-c2c-proposed.svg",
           "Figure 3 — Scheduled C2C Architecture: 1 Pod vs 2 Pods", width_inches=6.3)

add_horizontal_rule(doc)

# ===================================================================
# 8. END-TO-END FLOW
# ===================================================================
add_heading_with_color(doc, "8. End-to-End Flow", level=1)

add_heading_with_color(doc, "8.1 Flow Creation", level=2)
add_code_block(doc, '''1. Customer creates flow via CDC API:
   - source_type: SOURCE_SCHEDULE
   - source_data_types: []
   - script_schedule: "*/5 * * * *"
   - destinations: [{ type: DESTINATION_APPLICATION,
                      application: { type: PYTHON_SCRIPT, ... } }]

2. cdc.flow.api.service:
   - validateC2CService() → PASSES (SOURCE_SCHEDULE no longer blocked)
   - getSrcSvcConf()      → returns nil (scheduled flow detected)
   - NewCDCFlowHandler()  → CDCFlowHandler{InSvc: nil, OutSvc: soarLightSvc}
   - CreateFlow()         → skips InSvc, creates OutSvc ConfigMap + StatefulSet

3. Kubernetes creates:
   - ConfigMap:    cdc-soar-light-cm-{acctId}-{flowId}   (contains FLOW_JSON)
   - StatefulSet:  cdc-soar-light-service-{acctId}-{flowId}   (1 replica)

4. soar-light pod starts:
   - Reads FLOW_JSON from env var
   - cron.Start() fires the robfig/cron ticker goroutine (Change 7)
   - processEvent() → validateDestinations() → runDataFlow()
   - Detects: DataTypes=0, Schedule="*/5 * * * *"
   - Calls ProcessScheduledScript() → creates venv → registers cron job
   - Cron ticker invokes IScheduler.schedule() every 5 minutes''')

add_heading_with_color(doc, "8.2 Flow Deletion", level=2)
add_code_block(doc, '''1. Customer deletes flow via CDC API
2. cdc.flow.api.service:
   - DeleteFlow() → skips InSvc (nil), deletes OutSvc ConfigMap + StatefulSet
3. Kubernetes removes the soar-light pod''')

add_heading_with_color(doc, "8.3 Sync Reconciliation", level=2)
add_code_block(doc, '''1. syncCDCFlows() runs periodically
2. For scheduled flows:
   - Expects 1 ConfigMap (not 2)
   - Skips InSvc ConfigMap/StatefulSet reconciliation
   - Reconciles OutSvc ConfigMap + StatefulSet normally''')

add_horizontal_rule(doc)

# ===================================================================
# 9. IMPACT SUMMARY
# ===================================================================
add_heading_with_color(doc, "9. Impact Summary", level=1)

styled_table(doc,
    ["Component", "Impact", "Changes"],
    [
        ["cdc.flow.api.service/.../flow_handler_service.go", "Modified", "Remove SourceScheduleType from blocked list"],
        ["cdc.flow.api.service/.../flow_handler.go",         "Modified", "isScheduledFlow() helper; nil InSvc in constructor + CRUD"],
        ["cdc.flow.api.service/.../utils.go",                "Modified", "getSrcSvcConf() returns nil; sync handles nil InSvc"],
        ["soar-light/cmd/cloud/main.go",                     "Modified", "Add cron.Start() so scheduled jobs fire"],
        ["cdc.grpc-in/*",                                    "None",     "Not involved in scheduled flows"],
        ["CDC API contract",                                 "None",     "No API changes"],
        ["On-prem flows",                                    "None",     "Unaffected"],
        ["Existing event-driven C2C flows",                  "None",     "Behavior unchanged"],
    ],
    col_widths=[2.8, 0.8, 2.8]
)

add_horizontal_rule(doc)

# ===================================================================
# 10. TEST VALIDATION
# ===================================================================
add_heading_with_color(doc, "10. Test Validation", level=1)

add_heading_with_color(doc, "10.1 Build Verification", level=2)
add_code_block(doc, '''$ cd cdc.flow.api.service && go build ./pkg/svc/
# Clean — no errors

$ cd soar-light && go build ./cmd/cloud/
# Clean — no errors''')

add_heading_with_color(doc, "10.2 Unit Tests", level=2)
add_code_block(doc, '''$ go test ./pkg/svc/ -run "TestNewFlowConfig|Test_extractFlowId|Test_getStsName|Test_getCmName|Test_mergeJson"
# All PASS''')

add_heading_with_color(doc, "10.3 Live Environment Verification", level=2)
doc.add_paragraph("Flow 3650 created with SOURCE_SCHEDULE in the dev cluster:")
add_code_block(doc, '''$ kubectl get all -n cdc-data-flow | grep 3650
pod/cdc-soar-light-service-3006869-3650-0       1/1   Running
configmap/cdc-soar-light-cm-3006869-3650
statefulset.apps/cdc-soar-light-service-3006869-3650''')
p = doc.add_paragraph()
run = p.add_run("Result: 1 pod, 1 ConfigMap, 1 StatefulSet — no grpc-in resources created.")
run.bold = True

add_horizontal_rule(doc)

# ===================================================================
# 11. RESOLVED ISSUES
# ===================================================================
add_heading_with_color(doc, "11. Resolved Issues", level=1)

add_heading_with_color(doc, '11.1  soar-light: "destination config is empty"', level=2)
p = doc.add_paragraph()
run = p.add_run("Status: ")
run.bold = True
p.add_run("Resolved")

doc.add_paragraph(
    "The deployed soar-light Docker image was running an older code version (not release/v0.4.3). "
    "The error message format differed from the source code."
)
doc.add_paragraph(
    "Resolution: Redeployed with the correct release/v0.4.3 base image. Flow creation, cron "
    "registration, and health monitor all started working."
)

add_heading_with_color(doc, "11.2  Cron schedule() never fires in cloud mode", level=2)
p = doc.add_paragraph()
run = p.add_run("Status: ")
run.bold = True
p.add_run("Resolved (Change 7)")

doc.add_paragraph(
    "After fixing issue 11.1, the health monitor ran every 5 minutes but IScheduler.schedule() "
    "was never called. Root cause: cron.Start() was present in cmd/onprem/main.go but missing "
    "from cmd/cloud/main.go. Without it, robfig/cron registered jobs but never started its "
    "internal ticker goroutine."
)
doc.add_paragraph(
    "Resolution: Added svc.GetCron().Start() + defer cronScheduler.Stop() in "
    "cmd/cloud/main.go (Change 7)."
)

add_heading_with_color(doc, "11.3  Python Script Interface", level=2)
doc.add_paragraph(
    "Scheduled flows require only IScheduler — the IEventHandler interface is not used. "
    "The soar-light driver (infoblox_driver.py) detects the correct class via the "
    "--script-schedule flag automatically."
)

add_horizontal_rule(doc)

# ===================================================================
# 12. ARCHITECTURE DIAGRAMS
# ===================================================================
add_heading_with_color(doc, "12. Architecture Diagrams", level=1)
doc.add_paragraph(
    "The following diagrams provide visual context for the architecture and changes "
    "described in this document. All diagrams are embedded below as high-resolution images."
)

# List of diagrams to embed in section 12
diagram_entries = [
    ("architecture-and-data-flow.svg",           "Figure A — High-Level System Architecture (Cloud vs On-Prem)"),
    ("c2c-current-vs-scheduled-comparison.svg",  "Figure B — Current Event-Driven vs Proposed Scheduled C2C Flow"),
    ("c2c-cloud-flow.svg",                       "Figure C — C2C Data Flow: Dapr → Venv → Kafka → Python"),
    ("scheduled-c2c-proposed.svg",               "Figure D — Scheduled C2C Architecture (1 Pod vs 2 Pods)"),
    ("scheduled-cron-execution-lifecycle.svg",   "Figure E — Cron Execution Lifecycle: Pod Startup → cron.Start() → schedule()"),
    ("event-vs-scheduled.svg",                   "Figure F — Event Handler vs Scheduled Script Comparison"),
    ("script-execution-engine.svg",              "Figure G — Script Engine: Interfaces, Venv Lifecycle, Process Types"),
    ("kafka-topics-and-grpc-in-integration.svg", "Figure H — Kafka Topic Naming and gRPC-in Bridge"),
]

for svg_file, caption in diagram_entries:
    svg_path = os.path.join(SCRIPT_DIR, svg_file)
    if os.path.exists(svg_path):
        doc.add_page_break()
        add_diagram(doc, svg_file, caption, width_inches=6.5)
    else:
        doc.add_paragraph(f"[Diagram not found: {svg_file}]")

# Also include a reference table for all SVG files on disk
doc.add_page_break()
add_heading_with_color(doc, "Diagram Reference Index", level=2)
styled_table(doc,
    ["#", "Filename", "Description"],
    [
        ["1",  "architecture-and-data-flow.svg",            "High-level system architecture (Cloud vs On-Prem)"],
        ["2",  "c2c-current-vs-scheduled-comparison.svg",   "Current event-driven vs proposed scheduled C2C flow"],
        ["3",  "c2c-cloud-flow.svg",                        "C2C data flow: Dapr → Venv → Kafka → Python"],
        ["4",  "c2c-example-walkthrough.svg",               "C2C example walkthrough (Flow #43310)"],
        ["5",  "onprem-grpc-in-flow.svg",                   "On-Prem data flow: File config → gRPC → Python"],
        ["6",  "onprem-example-walkthrough.svg",            "On-Prem example walkthrough (Flow #23809)"],
        ["7",  "script-execution-engine.svg",               "Script engine: interfaces, venv lifecycle, process types"],
        ["8",  "kafka-topics-and-grpc-in-integration.svg",  "Kafka topic naming and gRPC-in bridge"],
        ["9",  "beginner-c2c-flow.svg",                     "C2C flow overview"],
        ["10", "beginner-onprem-flow.svg",                  "On-Prem flow overview"],
        ["11", "event-vs-scheduled.svg",                    "Event handler vs Scheduled script comparison"],
        ["12", "scheduled-c2c-proposed.svg",                "Scheduled C2C architecture (1 pod vs 2 pods)"],
        ["13", "scheduled-cron-execution-lifecycle.svg",    "Cron execution lifecycle: pod startup → cron.Start() → schedule()"],
    ],
    col_widths=[0.4, 3.2, 2.8]
)

# ===================================================================
# FOOTER — page numbers
# ===================================================================
footer = section.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run("Infoblox — Internal / Engineering  |  SOAR-Light C2C Scheduled Source Requirements  |  Page ")
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
# Page number field
fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
fp._p.append(fldChar1)
instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
fp._p.append(instrText)
fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
fp._p.append(fldChar2)

# ===================================================================
# SAVE
# ===================================================================
output_dir = os.path.dirname(os.path.abspath(__file__))
output_path = os.path.join(output_dir, "SOAR-Light_C2C_Scheduled_Source_Requirements.docx")
doc.save(output_path)
# Clean up PNG cache
import shutil
if os.path.exists(PNG_DIR):
    shutil.rmtree(PNG_DIR)

print(f"\n✅  Document saved: {output_path}")
print(f"    Sections: 12 + Appendix Diagrams")
print(f"    Embedded diagrams: 8 (SVG→PNG @200 DPI)")
print(f"    Inline figures: 3 (in sections 1, 3, 7)")
